import base64
import hashlib
import io
import re
import sys
from datetime import timezone, datetime

from flask import Flask
import requests
import hmac
from docx import Document




app = Flask(__name__)

username=""
password=""



@app.route("/")
def webpage():
    body = {"username":username, "pwd":password}
    auth = requests.post("https://www.fler.cz/api/rest/user/auth", json=body, verify=False,
           allow_redirects=False)

    secret_key = auth.json()['secret_key']
    session_id = auth.json()['session_id']

    request_path = "/api/rest/seller/products/list"
    auth_string = calculate_auth_string(request_path, secret_key, session_id)

    headers = {"X-FLER-AUTHORIZATION": auth_string}
    image_size='m' # options are m,s,b (medium, small, big)
    url_args = "?fields=title,description,keywords_tag,photo_main,colors,keywords_mat,attr2&photo_main="+image_size

    #styl ??? attr2...
    #klicova slova
    #material
    #barvy

    product_list_response = api_get(headers, request_path, url_args)
    # print(str(product_list_response.text))
    export_docx(product_list_response,image_size)

    return "Here will be some content printed out eventually"


def export_docx(product_list_response,image_size):
    cleanr = re.compile('<.*?>')

    document = Document()
    grouped_by_title = {}
    for product in product_list_response.json():
        if product['title'] in grouped_by_title:
            grouped_by_title.get(product['title']).append(product)
        else:
            grouped_by_title[product['title']] = [product]


    for product_title in grouped_by_title:
        document.add_heading(product_title,level=1)

        paragraph = document.add_paragraph()
        description = re.sub(cleanr, '', grouped_by_title[product_title][0]['description'])
        paragraph.add_run(description)

        document.add_paragraph().add_run("Varianty").bold = True
        for product in grouped_by_title[product_title]:
            paragraph = document.add_paragraph()
            image_url = product['photo_main'][image_size]
            image_response = requests.get(image_url, stream=True)
            image = io.BytesIO(image_response.content)
            paragraph.add_run().add_picture(image)
            paragraph.add_run('\nTags: ' + product['keywords_tag'])
            paragraph.add_run('\nColors: ' + product['colors'])

        document.add_page_break()


    document.save('demo.docx')
    # print(str(product_list_response.text))


def api_get(headers, request_path, url_args):
    return requests.get("https://www.fler.cz" + request_path + url_args, headers=headers)

def calculate_auth_string(request_path, secret_key, session_id):
    timestamp = str(int(datetime.utcnow().replace(tzinfo=timezone.utc).timestamp()))
    request_string = "GET\n" + timestamp + "\n" + request_path
    signature = hmac.new(secret_key.encode(), request_string.encode(), hashlib.sha1).hexdigest()
    auth_string = "API1_SESS" + " " + session_id + " " + timestamp + " " + base64.b64encode(signature.encode()).decode()
    return auth_string


if __name__ == "__main__":
    args = sys.argv[1:]
    username = args[0]
    password = args[1]
    webpage()
    # app.run('0.0.0.0',5002)


