import io
import re

import requests
from docx import Document
from docx.shared import Cm


from downloader.catalog_item_configuration import CustomizedCatalogItem
from downloader.image_utils import ImageUtils

NOT_SPECIFIED = "Nespecifikováno"


class DocxExporter:

    def __init__(self, stop_event):
        self.stop_event = stop_event

    def export_docx(self, product_list, image_size):
        print("Starting docx export.")
        cleanr = re.compile('<.*?>')
        customized_catalog_items = {}
        document = Document()
        grouped_by_title = {}
        for product in product_list:
            if product['title'] in grouped_by_title:
                grouped_by_title.get(product['title']).append(product)
            else:
                grouped_by_title[product['title']] = [product]

        for product_title in grouped_by_title:
            if (self.stop_event.is_set()):
                raise InterruptedError("Export interrupted.")

            document.add_heading(product_title, level=1)

            paragraph = document.add_paragraph()
            description = re.sub(cleanr, '', grouped_by_title[product_title][0]['description'])
            paragraph.add_run(description)

            short_description = re.sub(cleanr, '', grouped_by_title[product_title][0]['description_short'])
            paragraph = document.add_paragraph()
            paragraph.add_run("Kratky popis: ").bold = True
            paragraph.add_run(short_description)

            paragraph = document.add_paragraph()
            paragraph.add_run("Cena: ").bold = True
            paragraph.add_run(str(grouped_by_title[product_title][0]['price_without_prov']) + ' Kč (Fler: '
                           + str(grouped_by_title[product_title][0]['price']) + ' Kč)')


            self.print_category(grouped_by_title[product_title][0], document)

            document.add_paragraph().add_run("Varianty").bold = True

            variant_no = 1
            for product in grouped_by_title[product_title]:
                if (self.stop_event.is_set()):
                    raise InterruptedError("Export interrupted.")
                customized_catalog_items[product['id']] = self.write_variant(document, image_size, product, variant_no)
                variant_no += 1

            document.add_page_break()

        document.save('export.docx')
        print("Exported total of " + str(len(product_list)) + " records.")
        print("Docx export finished to file export.docx")
        return customized_catalog_items

    def print_category(self, product, document):
        # using custom category only (could be extended to use fler category (key category))
        paragraph = document.add_paragraph()
        paragraph.add_run("Kategorie: ").bold = True
        custom_category_id = product['sellcategory']
        if custom_category_id != '':
            paragraph.add_run(self.custom_categories[int(custom_category_id)])
        else:
            paragraph.add_run(NOT_SPECIFIED)

    def write_variant(self, document, image_size, product, variant_no):
        print("Writing: " + product['title'] + " " + str(product['id']))
        customized_write = False
        if product['id'] in self.custom_configurations:
            customized_catalog_item = self.custom_configurations[product['id']]
            customized_write = True
        else:
            customized_catalog_item = CustomizedCatalogItem(product['id'])
        customized_catalog_item.title = product['title']

        table = document.add_table(cols=2, rows=1, style='Table Grid')

        self.print_image(product, table.rows[0].cells[0], image_size)

        paragraph = table.rows[0].cells[1].paragraphs[0]
        if customized_write:
            paragraph.add_run("TYP " + str(variant_no) + ": ").bold = True
            paragraph.add_run(self.custom_configurations[product['id']].type)

        self.print_keywords(product, paragraph)
        self.print_material(product, paragraph)
        self.print_colors(product, paragraph)

        paragraph.add_run("\nKatalogové č.:\n").bold = True
        paragraph.add_run(str(product['id']))

        return customized_catalog_item

    def set_custom_category_list(self, custom_categories):
        self.custom_categories = custom_categories

    def set_custom_configurations(self, custom_configurations):
        self.custom_configurations = custom_configurations

    def print_image(self, product, table_cell, image_size):
        image_url = product['photo_main'][image_size]
        image_response = requests.get(image_url, stream=True)
        image = io.BytesIO(image_response.content)
        ImageUtils.save_image_to_tmp_folder(image, product['id'])
        paragraph = table_cell.paragraphs[0]
        table_cell.width = Cm(6.5)
        paragraph.add_run().add_picture(image, width=Cm(6.5))

    def print_keywords(self, product, paragraph):
        paragraph.add_run('\n\nKlíčová slova:\n').bold = True
        if product['keywords_tag'] is not None:
            paragraph.add_run(", ".join(product['keywords_tag'].split(",")))
        else:
            paragraph.add_run(NOT_SPECIFIED)

    def print_material(self, product, paragraph):
        # material is taken from keywords, but it could be taken from attr2 as well (keywords one is from custom worded tags, not predefined ones)
        paragraph.add_run('\nMateriál:\n').bold = True
        if product['keywords_mat'] is not None:
            paragraph.add_run(", ".join(product['keywords_mat'].split(",")))
        else:
            paragraph.add_run(NOT_SPECIFIED)

    def print_colors(self, product, paragraph):
        grouped_attr = product['attr2_grouped']
        if grouped_attr is not None:
            main_colors = list(filter(lambda element: element['group_type_ident'] == 'attr2_color', grouped_attr))
            paragraph.add_run('\nBarvy hlavní:\n').bold = True
            if len(main_colors) > 0:
                paragraph.add_run(",".join([item['ident'] for item in main_colors[0]['attr2']]))
            else:
                paragraph.add_run(NOT_SPECIFIED)
            secondary_colors = list(filter(lambda element: element['group_type_ident'] == 'attr2_color_secondary', grouped_attr))
            if len(secondary_colors) > 0:
                secondary_colors_str = ", ".join([item['ident'] for item in secondary_colors[0]['attr2']])
                if secondary_colors_str != "":
                    paragraph.add_run('\nBarvy vedlejší:\n').bold = True
                    paragraph.add_run(secondary_colors_str)



